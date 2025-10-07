# it convert our data into chunks and then chunks are stored in structured format in travel_records.json file

import os
import pickle
import json
import time
from docx import Document
from dotenv import load_dotenv
from openai import OpenAI

# === ENV SETUP ===
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=OPENAI_API_KEY)

# === Constants ===
FOLDER = "public"
CHUNKS_PICKLE_FILE = "chunks.pkl"
EXTRACTED_RECORDS_FILE = "travel_records.json"
MODEL_NAME = "gpt-4o-mini"  # Fast and cost-effective model
MAX_OUTPUT_TOKENS = 1000
DELAY_SECONDS = 1  # OpenAI has higher rate limits
BATCH_SIZE = 3

# === DOCX to Chunks ===
def extract_chunks_from_docx(file_path):
    doc = Document(file_path)
    chunks = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            chunks.append({
                "type": "paragraph",
                "full_text": text
            })

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                chunks.append({
                    "type": "table_row",
                    "full_text": " | ".join(row_text)
                })

    return chunks

# === Save Chunks ===
def save_chunks():
    all_chunks = []
    for filename in os.listdir(FOLDER):
        if filename.endswith(".docx"):
            path = os.path.join(FOLDER, filename)
            print(f"📄 Extracting chunks from: {filename}")
            all_chunks.extend(extract_chunks_from_docx(path))

    with open(CHUNKS_PICKLE_FILE, "wb") as f:
        pickle.dump(all_chunks, f)
    print(f"✅ Saved {len(all_chunks)} chunks to {CHUNKS_PICKLE_FILE}")

# === Load Chunks ===
def load_chunks():
    if not os.path.exists(CHUNKS_PICKLE_FILE):
        raise FileNotFoundError(f"❌ Missing: {CHUNKS_PICKLE_FILE}")
    with open(CHUNKS_PICKLE_FILE, "rb") as f:
        return pickle.load(f)

# === OpenAI API Call for 3-chunk batch ===
def openai_extract_batch(chunk_group):
    combined_input = "\n\n---\n\n".join([c["full_text"] for c in chunk_group])

    system_prompt = (
        "You are a data extraction engine. From the input text, extract all structured records present "
        "as valid JSON objects. Return them as a JSON array. Do not include explanations.\n\n"
        "Each input segment is separated by ---\n\n"
        "If no valid record is present in a section, skip it.\n\n"
        "Example output:\n"
        "[\n"
        "  {\n"
        "    \"type\": \"Travel Request\",\n"
        "    \"employee\": \"Shoaib Mkt\",\n"
        "    \"origin\": \"Haldwani\",\n"
        "    \"destination\": \"Bilaspur\",\n"
        "    \"date\": \"02-Apr-2025\",\n"
        "    \"email\": \"shoaib.mkt@shrirampistons.com\",\n"
        "    \"cost\": 1500\n"
        "  }\n"
        "]"
    )

    try:
        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": combined_input}
            ],
            temperature=0,
            max_tokens=MAX_OUTPUT_TOKENS
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"⚠️ OpenAI error: {e}")
        return None

# === Process All Chunks in Batches ===
def extract_all_records(chunks):
    print(f"🔍 Processing {len(chunks)} chunks in batches of {BATCH_SIZE}...")
    extracted = []

    for i in range(0, len(chunks), BATCH_SIZE):
        batch = chunks[i:i + BATCH_SIZE]
        print(f"🔄 Batch {i // BATCH_SIZE + 1}: Chunks {i + 1} to {i + len(batch)}")
        raw_output = openai_extract_batch(batch)
        time.sleep(DELAY_SECONDS)

        if raw_output is None:
            continue

        try:
            parsed = json.loads(raw_output)
            if isinstance(parsed, list):
                extracted.extend(parsed)
        except json.JSONDecodeError:
            print("❌ Skipped invalid JSON response.")
            continue

    return extracted

# === Save to JSON ===
def save_records(records):
    grouped = {}
    for r in records:
        rtype = r.get("type", "Unknown")
        grouped.setdefault(rtype, []).append(r)

    with open(EXTRACTED_RECORDS_FILE, "w", encoding="utf-8") as f:
        json.dump(grouped, f, ensure_ascii=False, indent=2)
    print(f"✅ Saved {len(records)} records to {EXTRACTED_RECORDS_FILE}")

# === Main Runner ===
def main():
    # Only extract and save chunks if not already saved
    if os.path.exists(CHUNKS_PICKLE_FILE):
        print(f"📦 Using existing chunks from {CHUNKS_PICKLE_FILE}")
    else:
        print("🔧 Chunks file not found. Generating new chunks...")
        save_chunks()

    # Load chunks from pickle
    chunks = load_chunks()

    # Process chunks through OpenAI to extract structured records
    records = extract_all_records(chunks)

    # Save final records to JSON
    save_records(records)

    print("🎉 All done! Use `travel_records.json` for structured access.")


if __name__ == "__main__":
    main()